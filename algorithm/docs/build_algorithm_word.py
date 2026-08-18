#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""生成《高光谱45项算法能力与服务封装方案》Word。"""

from __future__ import annotations

import argparse
import json
import re
import shutil
import textwrap
import time
from collections import Counter
from pathlib import Path
from typing import Any

import httpx
import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


DOCS_DIR = Path(__file__).resolve().parent
SOURCE_DIR = DOCS_DIR.parent / "source"
ASSET_DIR = DOCS_DIR / "word-assets"
OUTPUT_DOCX = DOCS_DIR / "高光谱45项算法能力与服务封装方案.docx"
LIST_MD = DOCS_DIR / "采集到算法-算法清单.md"
SERVICE_MD = DOCS_DIR / "当前服务简单介绍.md"
API_MD = DOCS_DIR / "算法API测试清单.md"

BLUE = "1F4E78"
LIGHT_BLUE = "D9EAF7"
GREEN = "2E7D32"
LIGHT_GREEN = "E2F0D9"
ORANGE = "C65911"
LIGHT_ORANGE = "FCE4D6"
GRAY = "666666"
LIGHT_GRAY = "F2F2F2"


# 每项算法的非公式化原理说明。详细用途、场景、输入和输出从权威清单解析。
PRINCIPLES = {
    1: "把测区边界转成平行航线，并结合航高、视场角、旁向/航向重叠率估算航带间距和触发点；本质是满足覆盖约束下的路径优化。",
    2: "以统一时钟或触发脉冲为基准，对各设备时间戳做偏移估计和插值匹配，建立“影像帧—GPS/IMU姿态”的一一对应关系。",
    3: "融合GPS位置与IMU角速度、加速度，通过组合导航滤波抑制单一传感器漂移，连续解算平台位置、速度和姿态。",
    4: "对帧序列做完整性、亮度分布、饱和比例、模糊度和POS连续性检查，把不可恢复的数据在进入高成本处理前拦截。",
    5: "利用云的高亮、低温或特定光谱响应，以及云影的暗值和空间邻接关系，输出像素级遮挡掩膜。",
    6: "用遮光条件下采集的暗参考估计探测器本底响应，并从每个像元的原始DN中扣除，降低固定模式偏置。",
    7: "依据坏像元标定表或统计异常检测定位坏点、坏列，再使用邻域、同波段或跨波段插值恢复连续影像。",
    8: "估计推扫方向各探测单元的系统性响应差异，通过列/行统计、频域滤波或低秩分解压制周期性条纹。",
    9: "依据实验室波长/空间标定模型，对不同视场位置的波长偏移和空间偏移进行重采样，保证同一波段含义一致。",
    10: "应用相机定标系数、积分时间和增益，将无量纲DN线性或非线性转换为有物理单位的辐亮度。",
    11: "利用多架次重叠区或伪不变地物建立景间转换关系，把不同光照与增益条件下的数据拉到同一辐射基准。",
    12: "用已知反射率参考板建立辐亮度到反射率的经验线关系，抵消当时光照条件，适合低空无人机快速定标。",
    13: "通过辐射传输模型估计大气吸收、散射和路径辐射，从传感器观测辐亮度反演地表真实反射率。",
    14: "建模太阳—地物—传感器几何导致的方向性反射差异，将不同观测角数据归一到统一参考角度。",
    15: "把每帧POS、相机内外参数代入共线方程或直接地理定位模型，计算像素在地面的初始坐标。",
    16: "结合相机模型、POS和DEM进行地形反投影与重采样，消除倾斜摄影和地形起伏造成的位置偏差。",
    17: "在航带重叠区提取并匹配同名点，优化相对位姿后进行几何融合，生成覆盖完整测区的连续数据立方体。",
    18: "在重叠区估计亮度/色彩差异，选择低冲突接缝线并进行渐变融合，减少拼接硬边和明暗跳变。",
    19: "通过控制点、特征匹配或地理坐标重投影，把高光谱、RGB和矢量边界统一到相同坐标系、分辨率和像元网格。",
    20: "按信噪比、水汽吸收区或统计异常剔除不可靠波段，并可结合平滑滤波降低随机噪声，为后续建模保留稳定信息。",
    21: "Savitzky–Golay在移动窗口内做低阶多项式拟合以平滑噪声；包络线去除则归一化光谱背景、突出吸收谷形态。",
    22: "按波段或样本计算均值方差、最小最大值，把不同量纲特征变换到可比较范围，避免大数值波段支配模型。",
    23: "PCA寻找最大方差正交方向，MNF同时考虑噪声协方差，ICA强调统计独立性；共同目标是压缩冗余波段。",
    24: "依据相关性、互信息、模型重要性或稀疏约束挑选少量高价值波段，减少计算量并降低过拟合。",
    25: "把光谱和空间上相近的邻接像元聚成均质对象，再以对象为分析单位，从而保留地物边界并减少椒盐噪声。",
    26: "围绕有标签像元截取固定窗口的空谱立方体，处理边界填充、标签对齐和训练/验证拆分，形成模型可消费张量。",
    27: "利用健康植被近红外高反射、红光强吸收的差异，计算(NIR-Red)/(NIR+Red)，得到标准化绿度指标。",
    28: "将红光替换为对叶绿素变化更敏感的红边波段，与近红外构造归一化差异，减轻密冠层NDVI饱和。",
    29: "EVI引入蓝光与增益项校正大气影响；SAVI/MSAVI引入土壤调节项，提升稀疏植被和裸土背景下的稳定性。",
    30: "通过近红外、短波红外或绿光波段的归一化差异，利用水分吸收和水体反射特性突出冠层含水或开放水体。",
    31: "在连续光谱的一阶导数或吸收特征上定位拐点、峰谷和面积，形成红边位置、谷深等可解释生理参数。",
    32: "用地面化验值监督建立光谱特征到连续生化指标的映射，PLS适合共线数据，随机森林和神经网络可拟合非线性。",
    33: "用PROSAIL等辐射传输模型模拟参数—光谱关系，再通过查找表、优化或贝叶斯方法反求LAI、叶绿素等物理参数。",
    34: "SVM寻找最大间隔分类边界，随机森林汇总多棵随机决策树投票；适合小样本光谱分类并可输出精度指标。",
    35: "把像素光谱与端元库向量计算光谱角或相似度，角度越小表示光谱形状越接近，从而完成已知物质识别。",
    36: "1D-CNN沿波长轴学习局部吸收模式，RNN建模波段序列依赖，在不使用空间邻域时完成像素级分类。",
    37: "2D-CNN学习空间纹理，3D-CNN同时在空间和光谱维卷积，利用邻域上下文提升相似地物的区分能力。",
    38: "Transformer用自注意力连接远距离波段或空间位置，GCN在超像素/样本图上传播信息，适合复杂全局关系。",
    39: "通过预训练、域适配、度量学习或少样本原型，把已有区域知识迁移到仅有少量标签的新区域。",
    40: "分割模型为每个像素预测目标概率，检测模型输出对象位置；当前示例以低NDVI得分、阈值和连通域提取胁迫斑块。",
    41: "把混合光谱表示为多个端元光谱的组合，约束丰度非负且和为一，估计一个像素内各物质所占比例。",
    42: "RX检测以背景均值和协方差建立统计分布，计算像素的马氏距离；距离异常大的像素被标记为未知异常。",
    43: "对已配准多时相数据做差异向量、指数变化或模型对比，经过阈值与分类后输出变化位置和变化类型。",
    44: "用形态学开闭运算、多数滤波、CRF或连通域面积规则清理孤立点和小斑块，使分类结果符合空间连续性。",
    45: "按地块矢量对分类图或连续指标做分区统计，将像素结果聚合为面积、均值、占比和阈值告警，形成决策级产品。",
}


def read_text(path: Path) -> str:
    return path.read_text(encoding="utf-8")


def parse_algorithms() -> list[dict[str, Any]]:
    """从三份文档解析45项算法，以最新API测试表作为状态权威来源。"""
    list_text = read_text(LIST_MD)
    api_text = read_text(API_MD)

    summaries: dict[int, tuple[str, str, str]] = {}
    for m in re.finditer(
        r"^\|\s*(\d{1,2})\s*\|\s*([^|]+?)\s*\|\s*([^|]+?)\s*\|\s*([^|]+?)\s*\|$",
        list_text,
        re.M,
    ):
        n = int(m.group(1))
        if 1 <= n <= 45 and n not in summaries:
            summaries[n] = tuple(x.strip().replace("**", "") for x in m.groups()[1:])

    details: dict[int, dict[str, str]] = {}
    headings = list(re.finditer(r"^####\s+(\d+)\.\s+(.+?)\s*$", list_text, re.M))
    for i, match in enumerate(headings):
        n = int(match.group(1))
        block = list_text[match.end() : headings[i + 1].start() if i + 1 < len(headings) else len(list_text)]
        fields = {}
        for key in ("作用", "使用场景", "数据输入", "数据输出"):
            fm = re.search(rf"\|\s*\*\*{key}\*\*\s*\|\s*(.*?)\s*\|", block)
            fields[key] = fm.group(1).strip().replace("**", "") if fm else ""
        details[n] = {"title": match.group(2).strip(), **fields}

    statuses: dict[int, dict[str, str]] = {}
    for m in re.finditer(
        r"^\|\s*(\d{2})\s*\|\s*`([^`]+)`\s*\|\s*(可运行|骨架)\s*\|\s*200\s*\|\s*True\s*\|\s*(True|False)\s*\|\s*(有|无)\s*\|$",
        api_text,
        re.M,
    ):
        n = int(m.group(1))
        statuses[n] = {
            "id": m.group(2),
            "status": m.group(3),
            "implemented": m.group(4),
            "files_status": m.group(5),
        }

    params: dict[int, str] = {}
    api_headings = list(re.finditer(r"^###\s+(\d+)\.\s+(.+?)\s*$", api_text, re.M))
    for i, match in enumerate(api_headings):
        n = int(match.group(1))
        block = api_text[match.end() : api_headings[i + 1].start() if i + 1 < len(api_headings) else len(api_text)]
        pm = re.search(r"-\s+\*\*params\*\*：`(.+?)`", block)
        params[n] = pm.group(1) if pm else "{}"

    items = []
    for n in range(1, 46):
        level, summary_title, summary = summaries[n]
        detail = details[n]
        state = statuses[n]
        items.append(
            {
                "number": n,
                "id": state["id"],
                "title": detail["title"] or summary_title,
                "level": level,
                "status": state["status"],
                "implemented": state["implemented"] == "True",
                "summary": summary,
                "purpose": detail["作用"],
                "scenario": detail["使用场景"],
                "input": detail["数据输入"],
                "output": detail["数据输出"],
                "params": params.get(n, "{}"),
                "principle": PRINCIPLES[n],
            }
        )
    return items


def packaging_strategy(item: dict[str, Any]) -> str:
    """生成与算法层级和实现状态一致的单项封装建议。"""
    level = item["level"]
    if level.startswith("L0") or "L0→L1" in level:
        mode = "采集后批处理或边缘端流水线节点"
        resource = "CPU优先；与相机SDK、定标文件和采集日志解耦"
    elif "L1→L2" in level or level == "L2":
        mode = "大文件异步任务"
        resource = "采用分块读取、保留空间参考和NoData；重投影/镶嵌任务需要独立工作目录"
    elif level == "L3":
        mode = "同步小样本演示、生产异步推理"
        resource = "模型类接口绑定模型版本与波段配置；GPU任务进入专用队列"
    else:
        mode = "业务聚合服务"
        resource = "输出稳定的地块ID、统计口径、专题图和告警规则版本"
    state = (
        "当前已有可运行实现，应补充参数校验、基准数据、性能指标和异常回滚后再进入生产。"
        if item["implemented"]
        else "当前仅完成统一接口骨架；必须接入经验证的核心算法、真实测试数据和精度验收，不能以HTTP 200替代能力完成。"
    )
    return f"建议作为{mode}封装。{resource}。统一返回任务标识、统计信息、产物URI和处理日志。{state}"


def validate_algorithms(items: list[dict[str, Any]]) -> None:
    numbers = [x["number"] for x in items]
    ids = [x["id"] for x in items]
    counts = Counter(x["status"] for x in items)
    assert numbers == list(range(1, 46)), f"编号异常：{numbers}"
    assert len(ids) == len(set(ids)) == 45, "算法ID重复或缺失"
    assert counts["可运行"] == 12 and counts["骨架"] == 33, f"状态统计异常：{counts}"
    for item in items:
        assert len(item["principle"]) >= 30, f"#{item['number']} 原理说明过短"
        strategy = packaging_strategy(item)
        assert len(strategy) >= 60, f"#{item['number']} 封装建议过短"
    print("算法数据校验通过：45 项；可运行 12；骨架 33")


def configure_chinese_font() -> None:
    candidates = [
        "/System/Library/Fonts/PingFang.ttc",
        "/System/Library/Fonts/STHeiti Light.ttc",
        "/System/Library/Fonts/Supplemental/Arial Unicode.ttf",
    ]
    for path in candidates:
        if Path(path).exists():
            from matplotlib import font_manager

            font_manager.fontManager.addfont(path)
            name = font_manager.FontProperties(fname=path).get_name()
            plt.rcParams["font.sans-serif"] = [name]
            break
    plt.rcParams["axes.unicode_minus"] = False


def save_figure(fig: plt.Figure, path: Path) -> Path:
    path.parent.mkdir(parents=True, exist_ok=True)
    fig.savefig(path, dpi=210, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    return path


def build_diagrams(asset_dir: Path) -> dict[str, Path]:
    """生成三张中文正式架构图。"""
    configure_chinese_font()
    asset_dir.mkdir(parents=True, exist_ok=True)
    assets: dict[str, Path] = {}

    # L0-L4 数据处理链
    fig, ax = plt.subplots(figsize=(12, 5.6))
    ax.axis("off")
    layers = [
        ("L0 原始采集", "DN帧 · POS · 日志\n#1–5", "#DCE6F1"),
        ("L1 传感器产品", "辐亮度 · 初始定位\n#6–11", "#D9EAD3"),
        ("L2 标准分析底图", "反射率正射Cube\n#12–26", "#FFF2CC"),
        ("L3 信息产品", "指数 · 分类 · 检测\n#27–43", "#FCE5CD"),
        ("L4 决策产品", "地块统计 · 告警\n#44–45", "#EADCF8"),
    ]
    for i, (title, desc, color) in enumerate(layers):
        x = 0.02 + i * 0.195
        rect = plt.Rectangle((x, 0.34), 0.165, 0.40, facecolor=color, edgecolor="#4D4D4D", lw=1.2)
        ax.add_patch(rect)
        ax.text(x + 0.0825, 0.64, title, ha="center", va="center", fontsize=13, weight="bold")
        ax.text(x + 0.0825, 0.47, desc, ha="center", va="center", fontsize=11, linespacing=1.6)
        if i < 4:
            ax.annotate("", xy=(x + 0.193, 0.54), xytext=(x + 0.168, 0.54), arrowprops={"arrowstyle": "->", "lw": 2, "color": "#1F4E78"})
    ax.text(0.5, 0.88, "高光谱数据产品 L0–L4 全链路与 45 项算法分布", ha="center", fontsize=18, weight="bold", color="#1F4E78")
    ax.text(0.5, 0.16, "原始计数 → 物理辐亮度 → 可分析的反射率立方体 → 信息提取 → 业务决策", ha="center", fontsize=12, color="#555555")
    assets["pipeline"] = save_figure(fig, asset_dir / "l0-l4-pipeline.png")

    # 服务封装架构
    fig, ax = plt.subplots(figsize=(12, 6.2))
    ax.axis("off")
    rows = [
        (0.82, "调用方", "App / GIS / Web / 第三方系统", "#DCE6F1"),
        (0.66, "API 网关", "FastAPI · 鉴权 · 限流 · OpenAPI · 版本路由", "#D9EAD3"),
        (0.50, "统一契约", "file / file2 / params → success / data / files / trace_id", "#FFF2CC"),
        (0.34, "算法适配器", "一算法一目录 · router.py 薄路由 · service.py 业务实现", "#FCE5CD"),
        (0.18, "执行与数据", "CPU/GPU 队列 · GeoTIFF/ENVI/GeoJSON/CSV · 对象存储", "#EADCF8"),
    ]
    for y, title, desc, color in rows:
        rect = plt.Rectangle((0.12, y - 0.055), 0.76, 0.11, facecolor=color, edgecolor="#555555", lw=1.1)
        ax.add_patch(rect)
        ax.text(0.20, y, title, ha="left", va="center", fontsize=13, weight="bold")
        ax.text(0.39, y, desc, ha="left", va="center", fontsize=11)
    for y1, y2 in zip([0.765, 0.605, 0.445, 0.285], [0.715, 0.555, 0.395, 0.235]):
        ax.annotate("", xy=(0.5, y2), xytext=(0.5, y1), arrowprops={"arrowstyle": "->", "lw": 1.8, "color": "#1F4E78"})
    ax.text(0.5, 0.95, "45 项算法统一服务封装架构", ha="center", fontsize=18, weight="bold", color="#1F4E78")
    ax.text(0.5, 0.06, "横向治理：配置中心 · 日志审计 · 指标监控 · 模型/算法版本 · 产物生命周期", ha="center", fontsize=11, color="#555555")
    assets["architecture"] = save_figure(fig, asset_dir / "service-architecture.png")

    # 请求调用链
    fig, ax = plt.subplots(figsize=(12, 5.6))
    ax.axis("off")
    nodes = [
        ("标准输入", "GeoTIFF\nGeoJSON / CSV"),
        ("统一请求", "POST /run\nfile + file2 + params"),
        ("算法执行", "校验 → 读取 → 计算\n→ 写出 → 统计"),
        ("标准响应", "JSON data\n+ files URI"),
        ("业务交付", "专题图 · 面积表\n告警 · API"),
    ]
    for i, (title, desc) in enumerate(nodes):
        x = 0.025 + i * 0.195
        rect = plt.Rectangle((x, 0.33), 0.16, 0.38, facecolor="#F7F9FB", edgecolor="#1F4E78", lw=1.4)
        ax.add_patch(rect)
        ax.text(x + 0.08, 0.62, title, ha="center", fontsize=13, weight="bold", color="#1F4E78")
        ax.text(x + 0.08, 0.45, desc, ha="center", fontsize=10.5, linespacing=1.5)
        if i < 4:
            ax.annotate("", xy=(x + 0.19, 0.52), xytext=(x + 0.163, 0.52), arrowprops={"arrowstyle": "->", "lw": 2, "color": "#C65911"})
    ax.text(0.5, 0.88, "一次算法调用的数据流与交付闭环", ha="center", fontsize=18, weight="bold", color="#1F4E78")
    assets["request_flow"] = save_figure(fig, asset_dir / "request-flow.png")
    return assets


def screenshot_pages(base_url: str, asset_dir: Path) -> dict[str, Path]:
    """使用真实服务页面采集 Swagger 与算法列表截图。"""
    from playwright.sync_api import sync_playwright

    assets: dict[str, Path] = {}
    with sync_playwright() as p:
        browser = p.chromium.launch(headless=True)
        page = browser.new_page(viewport={"width": 1600, "height": 1000}, device_scale_factor=1)
        page.goto(f"{base_url}/docs", wait_until="networkidle")
        time.sleep(1)
        path = asset_dir / "swagger-overview.png"
        page.screenshot(path=str(path), full_page=False)
        assets["swagger"] = path

        page.goto(f"{base_url}/api/v1/algorithms", wait_until="networkidle")
        path = asset_dir / "algorithm-list.png"
        page.screenshot(path=str(path), full_page=True)
        assets["algorithm_list"] = path
        browser.close()
    return assets


def render_result_panel(title: str, response: dict[str, Any], output: Path) -> Path:
    """把真实接口响应与真实PNG/TIF产物组合成报告截图。"""
    configure_chinese_font()
    files = response.get("files") or {}
    image_path: Path | None = None
    for value in files.values():
        candidate = Path(str(value))
        if candidate.suffix.lower() == ".png" and candidate.exists():
            image_path = candidate
            break
    fig, axes = plt.subplots(1, 2, figsize=(12, 5.3), gridspec_kw={"width_ratios": [1.35, 1]})
    ax_image, ax_text = axes
    if image_path:
        ax_image.imshow(plt.imread(image_path))
        ax_image.set_title(f"真实产物：{image_path.name}", fontsize=12)
        ax_image.axis("off")
    else:
        ax_image.axis("off")
        ax_image.text(0.5, 0.5, "该接口未返回PNG预览\n以真实JSON响应为证", ha="center", va="center", fontsize=14)
    ax_text.axis("off")
    safe = {
        "success": response.get("success"),
        "algorithm_id": response.get("algorithm_id"),
        "implemented": response.get("implemented"),
        "message": response.get("message"),
        "data": response.get("data"),
        "files": {k: Path(str(v)).name for k, v in files.items()},
    }
    text = json.dumps(safe, ensure_ascii=False, indent=2)
    ax_text.text(0, 1, "\n".join(textwrap.wrap(text, 48, replace_whitespace=False)), va="top", fontsize=8.8)
    fig.suptitle(title, fontsize=17, weight="bold", color="#1F4E78")
    fig.text(0.5, 0.02, "来源：本地 FastAPI 服务真实调用；截图生成时未包含密钥与绝对路径", ha="center", fontsize=9, color="#666666")
    return save_figure(fig, output)


def call_algorithm(base_url: str, algorithm_id: str) -> dict[str, Any]:
    testdata = SOURCE_DIR / "algorithms" / algorithm_id / "testdata"
    primary = next((p for p in sorted(testdata.iterdir()) if p.name.startswith("input.") and p.name != "input.npy"), None)
    if primary is None:
        raise FileNotFoundError(f"{algorithm_id} 缺少主输入")
    second = next((p for p in sorted(testdata.iterdir()) if p.name.startswith("file2.")), None)
    params_path = testdata / "params.json"
    params = json.loads(params_path.read_text(encoding="utf-8")) if params_path.exists() else {}
    handles = []
    try:
        f1 = primary.open("rb")
        handles.append(f1)
        files: dict[str, Any] = {"file": (primary.name, f1)}
        if second:
            f2 = second.open("rb")
            handles.append(f2)
            files["file2"] = (second.name, f2)
        with httpx.Client(timeout=180, trust_env=False) as client:
            response = client.post(f"{base_url}/api/v1/{algorithm_id}/run", files=files, data={"params": json.dumps(params)})
            response.raise_for_status()
            return response.json()
    finally:
        for handle in handles:
            handle.close()


def capture_service_assets(base_url: str, asset_dir: Path) -> dict[str, Path]:
    """采集真实页面和三个代表算法的结果截图。"""
    asset_dir.mkdir(parents=True, exist_ok=True)
    with httpx.Client(timeout=10, trust_env=False) as client:
        response = client.get(f"{base_url}/api/v1/algorithms")
        response.raise_for_status()
        catalog = response.json()
    (asset_dir / "algorithm-list-response.json").write_text(json.dumps(catalog, ensure_ascii=False, indent=2), encoding="utf-8")

    assets = screenshot_pages(base_url, asset_dir)
    examples = [
        ("27_ndvi", "NDVI 植被指数真实运行结果", "ndvi-result.png"),
        ("34_svm_rf_classify", "SVM / 随机森林分类真实运行结果", "classification-result.png"),
        ("42_anomaly_detect", "RX 光谱异常检测真实运行结果", "anomaly-result.png"),
    ]
    for algorithm_id, title, filename in examples:
        result = call_algorithm(base_url, algorithm_id)
        (asset_dir / f"{algorithm_id}-response.json").write_text(json.dumps(result, ensure_ascii=False, indent=2), encoding="utf-8")
        assets[algorithm_id] = render_result_panel(title, result, asset_dir / filename)
    print(f"真实服务截图完成：{len(assets)} 张")
    return assets


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_run_font(run, name: str = "等线", size: float | None = None, bold: bool | None = None, color: str | None = None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_field(paragraph, field_code: str, placeholder: str = "") -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = field_code
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = placeholder
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, text, end])


def add_heading(doc: Document, text: str, level: int) -> None:
    paragraph = doc.add_heading(text, level=level)
    paragraph.paragraph_format.keep_with_next = True


def add_key_value_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for key, value in rows:
        cells = table.add_row().cells
        cells[0].width = Cm(3.0)
        cells[1].width = Cm(13.5)
        cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cells[1].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cells[0], LIGHT_BLUE)
        r0 = cells[0].paragraphs[0].add_run(key)
        set_run_font(r0, bold=True, color=BLUE)
        r1 = cells[1].paragraphs[0].add_run(value)
        set_run_font(r1)


def add_picture(doc: Document, path: Path, caption: str, width: float = 6.4) -> None:
    if not path.exists():
        return
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(path), width=Inches(width))
    caption_p = doc.add_paragraph(caption)
    caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_p.style = doc.styles["Caption"]


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.3)
    section.right_margin = Cm(2.3)
    for style_name, font_name, size in [
        ("Normal", "等线", 10.5),
        ("Title", "方正小标宋简体", 26),
        ("Heading 1", "黑体", 18),
        ("Heading 2", "黑体", 15),
        ("Heading 3", "黑体", 12.5),
        ("Caption", "等线", 9),
    ]:
        style = doc.styles[style_name]
        style.font.name = font_name
        style._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
        style.font.size = Pt(size)
    doc.styles["Heading 1"].font.color.rgb = RGBColor.from_string(BLUE)
    doc.styles["Heading 2"].font.color.rgb = RGBColor.from_string(BLUE)
    normal = doc.styles["Normal"]
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.35

    for section in doc.sections:
        header = section.header.paragraphs[0]
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = header.add_run("高光谱 45 项算法能力与服务封装方案")
        set_run_font(run, size=9, color=GRAY)
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_field(footer, "PAGE", "1")


def add_cover(doc: Document) -> None:
    for _ in range(4):
        doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("高光谱 45 项算法能力\n与服务封装方案")
    set_run_font(run, "方正小标宋简体", 28, True, BLUE)
    doc.add_paragraph()
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("L0–L4 全链路能力地图 · FastAPI 统一服务 · 生产化路线")
    set_run_font(run, size=13, color=GRAY)
    for _ in range(8):
        doc.add_paragraph()
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run("版本：V1.0\n日期：2026年8月\n适用：领导汇报、产品规划、研发联调")
    set_run_font(run, size=11)
    doc.add_page_break()


def add_toc(doc: Document) -> None:
    add_heading(doc, "目录", 1)
    p = doc.add_paragraph()
    add_field(p, r'TOC \o "1-3" \h \z \u', "请在 Word 中右键更新目录")
    doc.add_page_break()


def add_overview_table(doc: Document, items: list[dict[str, Any]]) -> None:
    section = doc.add_section()
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.left_margin = section.right_margin = Cm(1.5)
    add_heading(doc, "45 项算法能力总览", 1)
    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["编号", "层级", "算法 ID", "算法/方法", "状态", "一句话用途"]
    for i, text in enumerate(headers):
        set_cell_shading(table.rows[0].cells[i], BLUE)
        run = table.rows[0].cells[i].paragraphs[0].add_run(text)
        set_run_font(run, size=9, bold=True, color="FFFFFF")
    set_repeat_table_header(table.rows[0])
    for item in items:
        cells = table.add_row().cells
        values = [f"{item['number']:02d}", item["level"], item["id"], item["title"], item["status"], item["summary"]]
        for i, value in enumerate(values):
            run = cells[i].paragraphs[0].add_run(str(value))
            set_run_font(run, size=8.2)
        set_cell_shading(cells[4], LIGHT_GREEN if item["implemented"] else LIGHT_ORANGE)
    portrait = doc.add_section()
    portrait.orientation = WD_ORIENT.PORTRAIT
    portrait.page_width = Cm(21.0)
    portrait.page_height = Cm(29.7)
    portrait.top_margin = Cm(2.2)
    portrait.bottom_margin = Cm(2.0)
    portrait.left_margin = Cm(2.3)
    portrait.right_margin = Cm(2.3)


def build_document(items: list[dict[str, Any]], assets: dict[str, Path], output_path: Path) -> Path:
    doc = Document()
    configure_document(doc)
    add_cover(doc)
    add_toc(doc)

    add_heading(doc, "一、执行摘要", 1)
    doc.add_paragraph(
        "本报告将高光谱数据从采集到业务决策划分为 L0–L4 五级产品，形成 45 项算法/工程能力清单，并以单进程 FastAPI 服务提供统一 HTTP 入口。"
        "根据 2026-08-09 最近一次自动测试记录，45/45 接口返回 HTTP 200 且 success=true；其中 12 项已有可运行实现并产出文件，33 项为统一契约骨架。"
    )
    p = doc.add_paragraph()
    run = p.add_run("重要口径：")
    set_run_font(run, bold=True, color=ORANGE)
    p.add_run("骨架接口可达不等于核心算法已实现；生产计划必须分别管理“接口覆盖率、算法实现率、精度达标率”。")
    add_key_value_table(
        doc,
        [
            ("能力覆盖", "45 项，覆盖航前规划、传感器校正、L2 标准产品、L3 信息提取和 L4 业务交付"),
            ("当前实现", "12 项可运行；33 项骨架"),
            ("接口验证", "HTTP 200 = 45/45；success=true = 45/45；可运行项产出 files = 12/12"),
            ("服务定位", "当前适合培训、联调与能力地图；生产化需补鉴权、异步调度、监控、算法精度验收"),
        ],
    )
    add_picture(doc, assets["pipeline"], "图 1  高光谱 L0–L4 全链路与 45 项算法分布")

    add_heading(doc, "二、服务现状与总体架构", 1)
    doc.add_paragraph(
        "当前服务采用“一个进程、一个端口、一算法一目录”的结构。调用方上传 GeoTIFF、GeoJSON 或 CSV，通过统一 `/api/v1/{algorithm_id}/run` 契约获取 JSON，"
        "其中 `data` 返回统计/指标，`files` 返回专题图、栅格、矢量或模型产物路径。该设计降低联调成本，但生产环境还需在执行、存储与治理层进一步拆分。"
    )
    add_picture(doc, assets["architecture"], "图 2  45 项算法统一服务封装架构")
    add_picture(doc, assets["request_flow"], "图 3  一次算法调用的数据流与交付闭环")

    add_overview_table(doc, items)

    add_heading(doc, "四、45 项算法逐项说明", 1)
    current_group = ""
    for item in items:
        group = item["level"].split("→")[0].replace("前", "")
        if group != current_group:
            current_group = group
            add_heading(doc, f"{current_group} 能力组", 2)
        add_heading(doc, f"{item['number']:02d}. {item['title']}", 3)
        status_text = "可运行（已有核心实现）" if item["implemented"] else "骨架（接口已预留，核心算法待实现）"
        add_key_value_table(
            doc,
            [
                ("算法标识", f"{item['id']} ｜ 层级：{item['level']} ｜ 状态：{status_text}"),
                ("一句话用途", item["summary"]),
                ("解决什么问题", item["purpose"]),
                ("原理简述", item["principle"]),
                ("典型场景", item["scenario"]),
                ("输入数据", item["input"]),
                ("输出结果", item["output"]),
                ("接口参数示例", item["params"]),
                ("封装与生产建议", packaging_strategy(item)),
            ],
        )
        doc.add_paragraph()

    add_heading(doc, "五、统一封装策略", 1)
    strategy_sections = [
        ("1. API 与版本层", "保留统一 `/api/v1/{algorithm_id}/run|health`，增加 OAuth2/JWT 或网关签名、租户隔离、限流、请求体大小限制和 API 版本冻结。"),
        ("2. 输入输出契约层", "通过算法元数据声明必填文件、允许格式、波段语义、坐标系、参数 Schema 和输出 Schema；上传后先做格式、CRS、波段数、NoData 和数值范围校验。"),
        ("3. 算法适配层", "坚持一算法一目录，router 只处理 HTTP，service 暴露稳定执行接口；厂商 SDK、遥感软件、Python 模型或外部命令均由 Adapter 屏蔽差异。"),
        ("4. 数据与产物层", "小文件可本地临时处理，生产统一接入对象存储；每个任务使用独立工作目录，产物带 trace_id、算法版本、输入摘要、CRS、生成时间和生命周期。"),
        ("5. 执行与资源层", "指数、统计等轻任务可同步；正射、镶嵌、深度学习等重任务进入异步队列。按 CPU/GPU、内存、预计时长打资源标签，支持超时、取消和幂等重试。"),
        ("6. 治理与可观测层", "记录参数、算法/模型版本、输入哈希、耗时、资源、状态和错误码；建设成功率、P95耗时、队列深度、GPU利用率、产物大小与精度漂移监控。"),
        ("7. 业务交付层", "算法原始栅格不是最终价值。统一把结果转成专题图、地块统计、面积占比、告警和可追溯报告，L4 服务负责业务口径而非重复实现底层算法。"),
    ]
    for heading, body in strategy_sections:
        add_heading(doc, heading, 2)
        doc.add_paragraph(body)

    add_heading(doc, "六、真实服务与算法产物截图", 1)
    screenshot_specs = [
        ("swagger", "图 4  FastAPI Swagger 交互文档（本地真实服务）", 6.3),
        ("algorithm_list", "图 5  `/api/v1/algorithms` 45 项算法列表（本地真实响应）", 6.3),
        ("27_ndvi", "图 6  NDVI 算法真实调用与产物", 6.3),
        ("34_svm_rf_classify", "图 7  SVM/随机森林分类真实调用与产物", 6.3),
        ("42_anomaly_detect", "图 8  RX 异常检测真实调用与产物", 6.3),
    ]
    for key, caption, width in screenshot_specs:
        if key in assets:
            add_picture(doc, assets[key], caption, width)

    add_heading(doc, "七、生产化路线建议", 1)
    add_key_value_table(
        doc,
        [
            ("第一阶段：加固", "对现有 12 项补参数 Schema、真实样本、精度基线、异常码、单元/回归测试和性能测试。"),
            ("第二阶段：高频补齐", "优先实现 L2 清洗/特征与 L3 指数、变化检测、后处理等业务高频算法，形成可组合流水线。"),
            ("第三阶段：重能力接入", "正射、大气校正、镶嵌等采用成熟库或厂商 SDK；深度模型进入 GPU 异步服务与模型注册中心。"),
            ("第四阶段：业务闭环", "围绕地块、作物和任务组织 L4 报表、告警、复核、反馈与模型迭代，形成可量化业务价值。"),
        ],
    )

    add_heading(doc, "八、附录：统一接口示例", 1)
    doc.add_paragraph("服务地址：http://127.0.0.1:28800　　交互文档：/docs　　算法列表：/api/v1/algorithms")
    code = (
        'curl -X POST "http://127.0.0.1:28800/api/v1/27_ndvi/run" \\\n'
        '  -F "file=@algorithms/27_ndvi/testdata/input.tif" \\\n'
        "  -F 'params={\"red_band\":2,\"nir_band\":3}'"
    )
    p = doc.add_paragraph()
    set_cell_shading(p._p.getparent(), LIGHT_GRAY) if False else None
    run = p.add_run(code)
    set_run_font(run, "Menlo", 9)
    doc.add_paragraph(
        "统一响应字段：success（调用是否完成）、algorithm_id（算法标识）、implemented（核心实现状态）、message（说明）、"
        "data（统计与指标）、files（产物路径/URI）。生产环境建议追加 trace_id、algorithm_version、duration_ms 与 warnings。"
    )

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    print(f"Word 已生成：{output_path}")
    return output_path


def verify_docx(path: Path, items: list[dict[str, Any]]) -> None:
    """回读并验证算法、图片关系及敏感信息。"""
    import zipfile

    doc = Document(path)
    text = "\n".join(p.text for p in doc.paragraphs)
    for item in items:
        heading = f"{item['number']:02d}. {item['title']}"
        assert text.count(heading) == 1, f"算法标题缺失或重复：{heading}"
    assert "LUMIR_API_KEY" not in text and "sk-" not in text, "疑似包含敏感密钥"
    with zipfile.ZipFile(path) as archive:
        media = [n for n in archive.namelist() if n.startswith("word/media/")]
        assert len(media) >= 8, f"图片数量不足：{len(media)}"
        assert "word/document.xml" in archive.namelist()
    print(f"Word 验证通过：45 项标题完整；内嵌图片 {len(media)} 张；未发现密钥")


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--validate-only", action="store_true")
    parser.add_argument("--diagrams-only", action="store_true")
    parser.add_argument("--capture", action="store_true")
    parser.add_argument("--base-url", default="http://127.0.0.1:28800")
    parser.add_argument("--build", action="store_true")
    args = parser.parse_args()

    items = parse_algorithms()
    validate_algorithms(items)
    if args.validate_only:
        return

    diagram_assets = build_diagrams(ASSET_DIR)
    if args.diagrams_only:
        print(f"架构图已生成：{ASSET_DIR}")
        return

    service_assets: dict[str, Path] = {}
    if args.capture:
        service_assets = capture_service_assets(args.base_url, ASSET_DIR)
    else:
        known = {
            "swagger": ASSET_DIR / "swagger-overview.png",
            "algorithm_list": ASSET_DIR / "algorithm-list.png",
            "27_ndvi": ASSET_DIR / "ndvi-result.png",
            "34_svm_rf_classify": ASSET_DIR / "classification-result.png",
            "42_anomaly_detect": ASSET_DIR / "anomaly-result.png",
        }
        service_assets = {key: path for key, path in known.items() if path.exists()}

    if args.build or args.capture:
        output = build_document(items, {**diagram_assets, **service_assets}, OUTPUT_DOCX)
        verify_docx(output, items)


if __name__ == "__main__":
    main()
